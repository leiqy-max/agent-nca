"""
DOCX to HTML Converter for Server-side Preview
Converts DOCX files to HTML without requiring client-side download
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from io import BytesIO

def docx_to_html(file_path):
    """
    Convert DOCX file to HTML with inline styles
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        HTML string with embedded styles
    """
    try:
        # 验证文件是否存在且可读
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 验证文件大小
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError("文件大小为 0，可能是空文件")
        
        # 尝试打开 DOCX 文件
        try:
            doc = Document(file_path)
        except Exception as e:
            # 如果是 ZIP 相关错误，说明文件可能损坏或不是有效的 DOCX
            if "zip" in str(e).lower() or "central directory" in str(e).lower():
                raise ValueError(f"文件格式错误：该文件不是有效的 DOCX 文件或文件已损坏")
            raise
        html_parts = []
        
        # HTML Header with styles
        html_parts.append("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        body {
            font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #f5f5f5;
        }
        .docx-container {
            background: white;
            padding: 60px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            border-radius: 4px;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #333;
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: bold;
        }
        h1 { font-size: 28px; border-bottom: 2px solid #e0e0e0; padding-bottom: 10px; }
        h2 { font-size: 24px; }
        h3 { font-size: 20px; }
        p {
            margin: 12px 0;
            color: #333;
            text-align: justify;
        }
        .center { text-align: center; }
        .right { text-align: right; }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
        .underline { text-decoration: underline; }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }
        table, th, td {
            border: 1px solid #ddd;
        }
        th, td {
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
        ul, ol {
            margin: 12px 0;
            padding-left: 40px;
        }
        li {
            margin: 6px 0;
        }
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
        }
        .watermark {
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 80px;
            color: rgba(0, 0, 0, 0.05);
            pointer-events: none;
            z-index: 1000;
            white-space: nowrap;
        }
    </style>
</head>
<body>
    <div class="docx-container">
""")
        
        # Process paragraphs
        for para in doc.paragraphs:
            para_html = process_paragraph(para)
            if para_html:
                html_parts.append(para_html)
        
        # Process tables
        for table in doc.tables:
            table_html = process_table(table)
            html_parts.append(table_html)
        
        # HTML Footer
        html_parts.append("""
    </div>
    <!-- <div class="watermark">仅供预览</div> -->
</body>
</html>
""")
        
        return ''.join(html_parts)
        
    except Exception as e:
        error_msg = str(e)
        # 提供更友好的错误提示
        if "zip" in error_msg.lower() or "central directory" in error_msg.lower():
            user_msg = "该文件不是有效的 DOCX 文件或文件已损坏"
        elif "FileNotFoundError" in str(type(e)):
            user_msg = "文件不存在"
        elif "文件格式错误" in error_msg:
            user_msg = error_msg
        else:
            user_msg = f"解析失败: {error_msg}"
        
        return f"""
<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>预览失败</title></head>
<body style="font-family: 'Microsoft YaHei', Arial; padding: 40px; text-align: center;">
    <div style="max-width: 600px; margin: 0 auto;">
        <h2 style="color: #e74c3c; margin-bottom: 20px;">📄 文档预览失败</h2>
        <div style="background: #fff3cd; border: 1px solid #ffc107; border-radius: 8px; padding: 20px; margin: 20px 0;">
            <p style="color: #856404; margin: 0; font-size: 16px;">{user_msg}</p>
        </div>
        <div style="color: #666; line-height: 1.8;">
            <p><strong>可能的原因：</strong></p>
            <ul style="text-align: left; display: inline-block;">
                <li>文件不是标准的 DOCX 格式</li>
                <li>文件在上传过程中损坏</li>
                <li>文件使用了不兼容的加密或保护</li>
            </ul>
        </div>
        <p style="color: #999; margin-top: 30px;">💡 建议下载后使用 Microsoft Office 或 WPS 打开查看</p>
    </div>
</body>
</html>
"""

def process_paragraph(para):
    """Process a single paragraph and return HTML"""
    if not para.text.strip():
        return '<p>&nbsp;</p>'
    
    # Determine paragraph style
    style_class = []
    alignment = ''
    
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        alignment = ' class="center"'
    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        alignment = ' class="right"'
    
    # Check if it's a heading
    if para.style.name.startswith('Heading'):
        level = para.style.name.replace('Heading ', '').strip()
        if level.isdigit() and 1 <= int(level) <= 6:
            return f'<h{level}{alignment}>{escape_html(para.text)}</h{level}>\n'
    
    # Process runs (text with formatting)
    text_parts = []
    for run in para.runs:
        text = escape_html(run.text)
        
        # Apply formatting
        if run.bold:
            text = f'<strong>{text}</strong>'
        if run.italic:
            text = f'<em>{text}</em>'
        if run.underline:
            text = f'<u>{text}</u>'
        
        text_parts.append(text)
    
    full_text = ''.join(text_parts) if text_parts else escape_html(para.text)
    
    return f'<p{alignment}>{full_text}</p>\n'

def process_table(table):
    """Process a table and return HTML"""
    html = ['<table>']
    
    for i, row in enumerate(table.rows):
        html.append('<tr>')
        for cell in row.cells:
            tag = 'th' if i == 0 else 'td'
            cell_text = escape_html(cell.text.strip())
            html.append(f'<{tag}>{cell_text}</{tag}>')
        html.append('</tr>')
    
    html.append('</table>\n')
    return ''.join(html)

def escape_html(text):
    """Escape HTML special characters"""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))

def get_file_preview_html(file_path, filename):
    """
    Get preview HTML for various file types
    
    Args:
        file_path: Path to file
        filename: Original filename
        
    Returns:
        HTML string or None if not supported
    """
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.docx':
        return docx_to_html(file_path)
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            escaped = escape_html(content)
            return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: "Courier New", monospace; padding: 40px; background: #f5f5f5; }}
        pre {{ background: white; padding: 30px; border-radius: 4px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); white-space: pre-wrap; word-wrap: break-word; }}
    </style>
</head>
<body>
    <pre>{escaped}</pre>
</body>
</html>
"""
        except Exception as e:
            return None
    
    return None
